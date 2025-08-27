from docx import Document

# Create a new Word Document
doc = Document()

# Add Title
doc.add_heading('Automation & Performance Engineer Resume', 0)

# Contact Info
doc.add_paragraph("""
[Your Full Name]
📍 [City, State, Country]    📞 [+Country Code-Phone Number]    📧 [you@example.com]
🔗 LinkedIn: [linkedin.com/in/yourprofile]    💻 GitHub: [github.com/yourusername]
""")

# Professional Summary
doc.add_heading('Professional Summary', level=1)
doc.add_paragraph("""Results-driven Automation and Performance Engineer with X+ years of experience...
""")

# Technical Skills
doc.add_heading('Technical Skills', level=1)
doc.add_paragraph("""
Automation Tools: Selenium WebDriver, TestNG, Cypress, RestAssured, Playwright
Performance Tools: JMeter, LoadRunner, Gatling, BlazeMeter
Languages: Java, Python, JavaScript, Bash
CI/CD Tools: Jenkins, GitLab CI, CircleCI, Azure DevOps
Version Control: Git, Bitbucket
Monitoring Tools: New Relic, Dynatrace, Grafana, Kibana
Bug Tracking: JIRA, Bugzilla
Cloud/Containers: AWS, Docker, Kubernetes
Databases: MySQL, PostgreSQL, MongoDB
Frameworks: BDD (Cucumber), TestNG, Pytest
Protocols: HTTP, REST, SOAP, WebSockets
""")

# Professional Experience
doc.add_heading('Professional Experience', level=1)
doc.add_paragraph("""Automation and Performance Engineer
Company Name – Location | Month YYYY – Present
- Designed and implemented scalable automated test frameworks using Selenium and TestNG...
""")
doc.add_paragraph("""Software QA Engineer
Previous Company – Location | Month YYYY – Month YYYY
- Automated regression tests for web and mobile applications using Appium and Python...
""")

# Project Highlights
doc.add_heading('Project Highlights', level=1)
doc.add_paragraph("""E-commerce Platform Optimization
- Led performance testing initiative for a high-traffic e-commerce site...
""")
doc.add_paragraph("""CI/CD Pipeline Automation
- Implemented parallel test execution and report generation in CI pipelines...
""")

# Certifications
doc.add_heading('Certifications', level=1)
doc.add_paragraph("""
- Certified JMeter Performance Tester – [Institute Name], [Year]
- Selenium WebDriver with Java – [Platform Name], [Year]
- AWS Certified Developer – Associate (optional)
""")

# Education
doc.add_heading('Education', level=1)
doc.add_paragraph("""Bachelor of Technology in Computer Science
[University Name], [Year of Graduation]
""")

# Optional Sections
doc.add_heading('Optional Sections', level=1)
doc.add_paragraph("""
- Publications / Blog / Open Source Contributions
- Languages (if multilingual)
- Awards / Achievements
""")

# Save the document
doc.save("Automation_Performance_Engineer_Resume.docx")
